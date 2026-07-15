# -*- coding: utf-8 -*-
"""Build manuscript_v7.docx from the v7 markdown source, embedding the updated
figures (PNG) and the supplementary material. Derived from build_docx.py (v6)
but pinned to the v7 source so the v6 build stays untouched."""
import re, os
from docx import Document
from docx.shared import Pt, RGBColor, Twips, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MAIN = r"C:/D/workbuddy/科研/虚拟敲除/gate1/rewiring_study/manuscript_v7.md"
SUPP = r"C:/D/workbuddy/科研/虚拟敲除/gate1/rewiring_study/Supplementary_v6.md"
OUT  = r"C:/D/workbuddy/科研/虚拟敲除/gate1/rewiring_study/manuscript_v7.docx"

CONTENT_W = 9026  # A4 with 1-inch margins, in DXA
FONT = "Times New Roman"
FIG_EMBED_SEEN = set()  # figures that had a markdown image embed (found or not)

doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21); sec.page_height = Cm(29.7)
sec.left_margin = Cm(2.54); sec.right_margin = Cm(2.54)
sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
doc.styles["Normal"].font.name = FONT
doc.styles["Normal"].font.size = Pt(11)

def set_run_font(run, size=11, bold=False, italic=False, sup=False, sub=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold; run.italic = italic
    run.font.superscript = sup; run.font.subscript = sub
    if color is not None:
        run.font.color.rgb = color

# ---------------- inline parser ----------------
TOKEN = re.compile(r"(<sup>.*?</sup>|<sub>.*?</sub>|\*\*\*[^*]*?\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*|\^(\d+))")

def parse_inline(text, bold=False, italic=False):
    runs = []
    pos = 0
    while pos < len(text):
        m = TOKEN.search(text, pos)
        if not m:
            if text[pos:]:
                runs.append(("text", text[pos:], bold, italic))
            break
        if m.start() > pos:
            runs.append(("text", text[pos:m.start()], bold, italic))
        tok = m.group(0)
        if tok.startswith("<sup>"):
            runs.append(("sup", tok[5:-6], bold, italic))
        elif tok.startswith("<sub>"):
            runs.append(("sub", tok[5:-6], bold, italic))
        elif tok.startswith("***"):
            runs.extend(parse_inline(tok[3:-3], bold=True, italic=True))
        elif tok.startswith("**"):
            runs.extend(parse_inline(tok[2:-2], bold=True, italic=italic))
        elif tok.startswith("*"):
            runs.extend(parse_inline(tok[1:-1], bold=bold, italic=True))
        else:  # ^N
            runs.append(("sup", tok[1:], bold, italic))
        pos = m.end()
    return runs

def add_runs(par, text, size=11):
    for item in parse_inline(text):
        kind, val = item[0], item[1]
        b = item[2] if len(item) > 2 else False
        it = item[3] if len(item) > 3 else False
        r = par.add_run(val)
        set_run_font(r, size, bold=(b or kind in ("bold", "bolditalic")),
                     italic=(it or kind in ("italic", "bolditalic")),
                     sup=(kind == "sup"), sub=(kind == "sub"))

# ---------------- headings ----------------
def add_title(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); set_run_font(r, 18, bold=True)

def add_heading(text, level):
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    r = p.add_run(text)
    set_run_font(r, {1: 14, 2: 12, 3: 11}[level], bold=True)

# ---------------- tables ----------------
def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def parse_table_row(line):
    s = line.strip()
    if s.startswith("|"): s = s[1:]
    if s.endswith("|"): s = s[:-1]
    return [c.strip() for c in s.split("|")]

def is_separator(line):
    s = line.strip()
    if s.startswith("|"): s = s[1:]
    if s.endswith("|"): s = s[:-1]
    cells = [c.strip() for c in s.split("|")]
    return len(cells) > 0 and all(re.match(r"^:?-+:?$", c) for c in cells)

def build_table(rows, size=9):
    ncols = len(rows[0])
    widths = [CONTENT_W // ncols] * ncols
    widths[0] += CONTENT_W - sum(widths)
    t = doc.add_table(rows=len(rows), cols=ncols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for i, row in enumerate(rows):
        for j, celltext in enumerate(row):
            c = t.cell(i, j)
            c.width = Twips(widths[j])
            c.paragraphs[0].text = ""
            para = c.paragraphs[0]
            add_runs(para, celltext, size=size)
            if i == 0:
                for r in para.runs:
                    r.bold = True
                shade_cell(c, "D9E2F3")
    doc.add_paragraph()  # spacing after table

def flush_table(buf):
    if not buf:
        return
    tbl_lines = [l for l in buf if l.strip().startswith("|")]
    rows = [parse_table_row(l) for l in tbl_lines]
    rows = [r for idx, r in enumerate(rows) if not is_separator(tbl_lines[idx])]
    if rows:
        build_table(rows)
    buf.clear()

# ---------------- figure placeholder ----------------
FIGDIR = os.path.join(os.path.dirname(MAIN), "figures")
FIG_MAP = {
    "1": "figure1_tsc_gnn_conceptual_framework_v4.png",
    "2": "figure2_prediction_benchmark.png",
    "3": "figure3_rewiring_heatmaps.png",
    "4": "figure4_cross_species.png",
    "5": "figure5_drug_reversal.png",
    "6": "figure6_L5_triangulation.png",
    "7": "figure7_evidence_ladder.png",
    "8": "cellchat_rewiring_heatmap.png",
    "S1": "figureS1_pc_correction.png",
}
def figure_png(num):
    fn = FIG_MAP.get(num)
    if fn:
        p = os.path.join(FIGDIR, fn)
        if os.path.exists(p):
            return p
    return None

def add_picture_centered(path, width_cm=15):
    doc.add_picture(path, width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_fig_placeholder(label):
    p = doc.add_paragraph()
    r = p.add_run(f"[Figure {label} — figure image to be inserted]")
    set_run_font(r, 10, italic=True, color=RGBColor(0x80, 0x80, 0x80))

# ---------------- generic line processor ----------------
def process_lines(lines, single_hash_style="Title", double_hash_style="Heading 1"):
    buf = []
    i = 0; n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("|"):
            buf.append(line); i += 1; continue
        flush_table(buf)
        # markdown image embed: ![alt](path)
        if stripped.startswith("!["):
            m = re.match(r"^!\[(.*?)\]\((.*?)\)\s*$", stripped)
            if m:
                alt, rel = m.group(1), m.group(2)
                fn = re.search(r"Figure\s+(\d+)", alt)
                num = fn.group(1) if fn else "?"
                FIG_EMBED_SEEN.add(num)
                full = os.path.join(os.path.dirname(MAIN), rel)
                if os.path.exists(full):
                    try:
                        doc.add_picture(full, width=Cm(14))
                    except Exception:
                        add_fig_placeholder(num)
                else:
                    p = doc.add_paragraph()
                    r = p.add_run(f"[Figure {num} — image file not found: {rel}]")
                    set_run_font(r, 10, italic=True, color=RGBColor(0x80, 0x80, 0x80))
                i += 1; continue
        if stripped == "" or stripped == "---":
            i += 1; continue
        if stripped.startswith("$$"):
            inner = stripped
            if not inner.endswith("$$"):
                eq = [inner]; i += 1
                while i < n and not lines[i].strip().endswith("$$"):
                    eq.append(lines[i].strip()); i += 1
                if i < n:
                    eq.append(lines[i].strip()); i += 1
                inner = " ".join(eq)
            inner = inner.strip("$").strip()
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(inner); set_run_font(r, 11, italic=True)
            i += 1; continue
        if stripped.startswith(">"):
            bq = []
            while i < n and lines[i].strip().startswith(">"):
                bq.append(lines[i].strip().lstrip(">").strip()); i += 1
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(0.4)
            add_runs(p, " ".join(bq))
            continue
        if stripped.startswith("- "):
            items = []
            while i < n and lines[i].strip().startswith("- "):
                items.append(lines[i].strip()[2:]); i += 1
            for it in items:
                if "Supplementary Figure" in it:
                    m = re.search(r"Supplementary Figure\s+(S\d+)", it)
                    if m:
                        png = figure_png(m.group(1))
                        if png:
                            try:
                                add_picture_centered(png, width_cm=11)
                            except Exception:
                                add_fig_placeholder(m.group(1))
                        else:
                            add_fig_placeholder(m.group(1))
                p = doc.add_paragraph(style="List Bullet")
                add_runs(p, it)
            continue
        if re.match(r"^\d+\.\s", stripped):
            items = []
            while i < n and re.match(r"^\d+\.\s", lines[i].strip()):
                items.append(re.sub(r"^\d+\.\s", "", lines[i].strip())); i += 1
            for it in items:
                p = doc.add_paragraph(style="List Number")
                add_runs(p, it)
            continue
        if stripped.startswith("#"):
            if stripped.startswith("#### "):
                add_heading(stripped[5:], 3)
            elif stripped.startswith("### "):
                add_heading(stripped[4:], 2)
            elif stripped.startswith("## "):
                lvl = 1 if double_hash_style == "Heading 1" else 2
                add_heading(stripped[3:], lvl)
            else:
                if single_hash_style == "Title":
                    add_title(stripped[2:].strip())
                else:
                    lvl = 1 if single_hash_style == "Heading 1" else 2
                    add_heading(stripped[2:].strip(), lvl)
            i += 1; continue
        # figure legend entry: embed the figure image above its legend
        m = re.match(r"\*\*Figure\s+(\d+)", stripped)
        if m:
            num = m.group(1)
            png = figure_png(num)
            if png and num not in FIG_EMBED_SEEN:
                try:
                    add_picture_centered(png, width_cm=15)
                except Exception:
                    add_fig_placeholder(num)
            elif png is None and num not in FIG_EMBED_SEEN:
                add_fig_placeholder(num)
            p = doc.add_paragraph()
            add_runs(p, stripped)
            i += 1
            continue
        # default: render any other line as a normal paragraph
        p = doc.add_paragraph()
        add_runs(p, stripped)
        i += 1
    flush_table(buf)

# ---------------- split into sections ----------------
def split_sections(lines):
    preamble = []
    out = []
    cur_head = None; cur_body = []
    for l in lines:
        if re.match(r"^## ", l):
            if cur_head is not None:
                out.append((cur_head, cur_body))
            cur_head = l[3:].strip(); cur_body = []
        else:
            if cur_head is None:
                preamble.append(l)
            else:
                cur_body.append(l)
    if cur_head is not None:
        out.append((cur_head, cur_body))
    return preamble, out

# ---------------- main assembly ----------------
main_text = open(MAIN, encoding="utf-8").read()
# idempotent safety: normalise any stray triple-asterisk author marker
main_text = main_text.replace("Wenwu Huang^1,***", "Wenwu Huang^1,**")
main_lines = main_text.splitlines()

preamble, sections = split_sections(main_lines)

# title
title_line = next(l for l in preamble if l.startswith("# "))
add_title(title_line[2:].strip())
# rest of preamble (authors, affiliations, blockquote)
rest_preamble = [l for l in preamble if not l.startswith("# ")]
process_lines(rest_preamble)

# body sections (render in document order; only References is deferred to the end)
refs_body = None
for head, body in sections:
    if head == "References":
        refs_body = body
    else:
        add_heading(head, 1)
        process_lines(body)

# NOTE: the supplementary material is now embedded inline in manuscript_v7.md
# (synced from the docx), so it is no longer appended from Supplementary_v6.md.

# references last (plain paragraphs, matching the docx layout)
if refs_body is not None:
    doc.add_page_break()
    add_heading("References", 1)
    for l in refs_body:
        s = l.strip()
        if not s:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(0.3)
        p.paragraph_format.first_line_indent = Pt(-0.3)
        add_runs(p, s)

# fix <w:zoom> missing required percent attribute (python-docx quirk)
settings = doc.settings.element
zoom = settings.find(qn("w:zoom"))
if zoom is None:
    zoom = OxmlElement("w:zoom")
    settings.append(zoom)
zoom.set(qn("w:percent"), "100")

doc.save(OUT)
print("Saved:", OUT)
print("Sections rendered:", len(sections), "| references deferred:", refs_body is not None)
