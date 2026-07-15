# -*- coding: utf-8 -*-
"""Translate manuscript_v4.md into a Chinese Word document using python-docx."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r"C:\D 盘\科研\虚拟敲除\gate1\rewiring_study"
IMG = os.path.join(BASE, "figures", "figure1_tsc_gnn_conceptual_framework_v4.png")
OUT = os.path.join(BASE, "manuscript_v4_中文版.docx")

doc = Document()

# ---------- font helpers ----------
def set_run_font(run, ea="宋体", latin="Times New Roman", size=None, bold=None, italic=None, color=None):
    run.font.name = latin
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), ea)
    rfonts.set(qn('w:ascii'), latin)
    rfonts.set(qn('w:hAnsi'), latin)

def style_set_font(style, ea, latin, size, bold=False):
    style.font.name = latin
    style.font.size = Pt(size)
    style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), ea)
    rfonts.set(qn('w:ascii'), latin)
    rfonts.set(qn('w:hAnsi'), latin)

# Normal
style_set_font(doc.styles['Normal'], "宋体", "Times New Roman", 11)
# Headings
style_set_font(doc.styles['Heading 1'], "黑体", "Times New Roman", 15, bold=True)
style_set_font(doc.styles['Heading 2'], "黑体", "Times New Roman", 12.5, bold=True)
style_set_font(doc.styles['Heading 3'], "黑体", "Times New Roman", 11.5, bold=True)

# ---------- low-level builders ----------
def P(text="", bold=False, italic=False, size=None, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, italic=italic)
    return p

def rich(p, text):
    """Parse **bold** and *italic* markers into runs within paragraph p."""
    parts = text.split('**')
    for i, part in enumerate(parts):
        if part == '':
            continue
        if i % 2 == 1:
            sub = part.split('*')
            for j, s in enumerate(sub):
                if s == '':
                    continue
                r = p.add_run(s)
                r.bold = True
                set_run_font(r, italic=(j % 2 == 1))
        else:
            sub = part.split('*')
            for j, s in enumerate(sub):
                if s == '':
                    continue
                r = p.add_run(s)
                set_run_font(r, italic=(j % 2 == 1))

def P_rich(text, size=None, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    rich(p, text)
    return p

def H(text, level):
    if level == 1:
        p = doc.add_paragraph(style='Heading 1')
    elif level == 2:
        p = doc.add_paragraph(style='Heading 2')
    else:
        p = doc.add_paragraph(style='Heading 3')
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    return p

def B(text, size=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_run_font(r, size=size)
    return p

def NUM(text, size=None):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_run_font(r, size=size)
    return p

def MATH(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, latin="Cambria Math", ea="宋体", size=11, italic=True)
    return p

def NOTE(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(text)
    set_run_font(r, italic=True, size=10, color=RGBColor(0x55, 0x55, 0x55))
    return p

def add_table(headers, rows, widths=None, font_size=9):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        rp = hdr[i].paragraphs[0].add_run(h)
        set_run_font(rp, bold=True, size=font_size)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            rp = cells[i].paragraphs[0].add_run(str(val))
            set_run_font(rp, size=font_size)
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows:
                r.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

def HR():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pbdr.append(bottom)
    pPr.append(pbdr)
    p.paragraph_format.space_after = Pt(6)

# =====================================================================
# TITLE BLOCK
# =====================================================================
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("恢复时序性调控重连：一种应用于缺血性卒中的可解释、基于图的虚拟扰动框架")
set_run_font(r, ea="黑体", latin="Times New Roman", size=17, bold=True)
sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sp.add_run("Recovering Temporal Regulatory Rewiring: An Interpretable Graph-Based Virtual Perturbation Framework Applied to Ischemic Stroke")
set_run_font(r, ea="宋体", latin="Times New Roman", size=10, italic=True, color=RGBColor(0x66,0x66,0x66))

# Authors
ap = doc.add_paragraph(); ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ap.add_run("Ming Luo¹, Yimin Mei¹, Wangyang Ye¹, Wenwu Huang¹,*")
set_run_font(r, size=11)
bp = doc.add_paragraph(); bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = bp.add_run("¹ 温州医科大学附属第五医院（丽水市中心医院）神经外科，中国浙江丽水")
set_run_font(r, size=10)
cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cp.add_run("* Ming Luo 与 Yimin Mei 对本工作贡献相同。")
set_run_font(r, size=10, italic=True)
dp = doc.add_paragraph(); dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = dp.add_run("* 通讯作者：黄文武（Wenwu Huang），温州医科大学附属第五医院（丽水市中心医院）神经外科，中国浙江丽水。Email: hwenwu321@gmail.com")
set_run_font(r, size=10, italic=True)

NOTE("目标期刊：Nature Methods / Patterns (Cell Press) / Nature Biotechnology / Bioinformatics　|　稿件类型：带验证的方法学论文　|　版本：v5 — 2026-07-10")

# =====================================================================
# TERMINOLOGY LEDGER
# =====================================================================
H("术语总表", 1)
NOTE("读者提示：本文使用若干专业缩写。下表术语总表（Terminology Ledger）在首次出现时给出定义。正文尽量减少缩写堆叠：任何句子包含的陌生缩写不超过两个。章节标号（如 §3.5）尽量精简。")
add_table(
    ["规范术语", "定义", "首次出现展开"],
    [
        ["TSC-GNN", "时序与细胞状态条件化的图神经网络", "先完整写出，后简用 “TSC-GNN”"],
        ["ΔW", "边级重连（跨越某转换的耦合变化）", "在式(1)中定义"],
        ["GRN", "基因调控网络", "GSEA/DoRothEA 有向因果图"],
        ["DoRothEA", "文献整理的 TF→靶基因调控子数据库", "Alvarez-Garcia 2019, Nat Commun"],
        ["MCAO", "大脑中动脉闭塞", "小鼠卒中模型"],
        ["pseudotime", "连续轨迹（BEAM；Barry 2022）", "注：非 velocity"],
        ["A+C 模式", "跨模态分析：A=SigCom LINCS 特征匹配，C=Replogle K562 单细胞 CRISPRi 调控子响应", "在 §3.8 定义"],
        ["L1–L5", "五级证据阶梯", "表 1"],
        ["emp p", "经验置换 p 值", "在 n_perm 百万量级上计算"],
    ],
    widths=[1.3, 4.0, 2.0], font_size=9.5
)

# =====================================================================
# ABSTRACT
# =====================================================================
H("摘要（约 220 词）", 1)
P_rich("**背景。** 基因调控网络（GRN）的时序性重连驱动卒中后的组织修复，但现有的虚拟扰动方法要么缺乏时间分辨率、忽略细胞间通讯，要么不提供可解释的边级输出，以机制换取了准确性。")
P_rich("**结果。** 我们提出 TSC-GNN，一个从多时间点单细胞转录组中恢复边级 GRN 重连、主调控因子动态以及药物重定位假设的框架。在缺血性卒中（两个小鼠队列，24 h → 14 d）中，它恢复了典型的修复程序——急性炎症 → 少突胶质细胞髓鞘再生（Sox10→Plp1，ΔW = +0.51）→ 恢复——并具有跨队列可重复性（ρ = 0.48–0.55，p < 10⁻¹⁵）。恢复出的程序在人类 GRN 中保守（SOX10→髓鞘 OR = 27），并在人类卒中血液中被激活（n = 39，BH-q = 10⁻⁸）。损伤特征映射到神经保护剂（HDAC 抑制剂、他汀类药物）。通过三角验证三种公开扰动模态（谱系内敲除、L1000 过表达、K562 单细胞 CRISPRi）表明，程序级因果支持是“情境门控”的：在生物学恰当的扰动中恢复，在脱离情境的癌细胞系中缺失——这一梯度正是“可解释性优先”框架所预测的。")
P_rich("**边界。** 在固定因果图 + 线性读头的设定下，图结构并未将扰动预测提升到线性基线之上（90/90 配置中 0 例提升），我们公开报告这一结果。TSC-GNN 的贡献在于“可解释性”，而非准确性。")
P_rich("**可用性。** 所有代码与分析流程均以可复现脚本形式提供（conda，SHA-256 清单）。公开数据：GSE174574、GSE225948、GSE16561、GSE269122、GSE273163；Figshare（Replogle 2022，K562 伪批量）。")

# =====================================================================
# INTRODUCTION
# =====================================================================
H("引言", 1)
P_rich("缺血性卒中触发一系列复杂、随时间协调展开的分子事件级联——从急性兴奋毒性和神经炎症，经亚急性修复启动，到慢性重塑[1, 2]。单细胞转录组学已开始解析驱动各阶段的细胞类型与基因程序[3, 4]，然而将转录因子（TF）与其靶基因跨越时间耦合起来的“调控逻辑”仍鲜有描绘。知道在卒中时间进程中哪些 TF→靶连接增强或减弱——即“边级重连”——是理解修复如何被协调、以及为何在部分患者中最终失败所必需的。")
P_rich("若干计算框架已处理相关问题。**CellOracle**[5] 利用已有 GRN 并在计算机中扰动 TF 表达以预测下游转录组变化；**SCENIC** 与 **SCENIC+**[6, 7] 从共表达推断调控子活性；**NicheNet**[8] 优先排序驱动靶基因变化的配体–受体链接；**GEARS**[9] 利用图神经网络（GNN）预测 Perturb-seq 结果。这些方法推动了领域发展，但各自在时序重连方面存在范围局限：它们或仅在单时间点运行、或将 GRN 重连与细胞间通讯割裂、或将扰动结果压缩为单一准确性分数，而不揭示“哪些”边发生了改变以及“为何”改变。")
P_rich("近期大规模基础模型——**scGPT**[10]、**scFoundation**[11]——通过从数百万细胞学习通用表征，达到了最先进的细胞状态预测水平。然而其可解释性仍然有限：调控图要么缺失，要么隐式编码在注意力权重中，难以作为有向、因果的重连加以分析。反之，诸如 **RegVelo**[12]、**TemporalVAE**[13]、**RENGE**[14] 等时序方法虽建模动态基因调控，却聚焦于速度（velocity）或轨迹推断，而非边级因果图重连。因此，我们识别出的缺口并非缺乏预测模型，而是缺乏“可解释、边级、时序性的 GRN 重连”——它能整合 TF–靶耦合与细胞间通讯，并生成可检验、可转化的输出。")
P_rich("图神经网络（GNN）[15, 16, 17] 是在图基板上建模扰动的天然选择：它们沿已知边传播扰动信号并产生新的隐空间嵌入[18]。这一能力带有一种鲜被审视的注意事项。在固定、基于先验知识的 GRN 上运行的 GNN，只有在将图结构视为因果、且读头本身设计为可解释时，才能生成边级重连摘要——即每个 TF→靶耦合在某一生物转换中如何变化。若仅以预测准确性来评价同一 GNN，则可解释性优势不可见，该方法便退化为一个浅层扰动预测器，无法超越在相同特征空间上运行的线性基线[19]。**因此，我们将 GNN 用于扰动建模的两种用途区分开来：以准确性为导向（预测结果）与以可解释性为导向（揭示调控变化）。TSC-GNN 追求后者。**")
P_rich("**TSC-GNN** 是一个以可解释性而非准确性构建的、时序与细胞状态条件化的 GNN 框架。给定多时间点、多条件的 scRNA-seq 数据集，以及一份文献整理的有向因果 GRN（DoRothEA [20]），它（i）构建时间–状态条件化图，（ii）沿该图传播虚拟扰动以获得重连后的隐空间嵌入，（iii）读出“边级重连”（ΔW，带置换显著性）、主调控因子排序、通过 CellChat[21] 实现的细胞间通讯重塑，以及通过 LINCS L1000[22] 实现的的药物重定位图谱。该框架沿五级证据阶梯（L1–L5，表 1；可视化概览见图 7）进行验证，从跨队列可重复性逐级上升到独立的公开扰动因果支持。")
P_rich("缺血性卒中是检验该框架的概念验证，而非其本身的生物学发现。它的作用是表明所恢复的调控程序是连贯、可重复、跨物种保守，并受到独立公开扰动数据的方向性支持。核心方法学贡献在于：固定因果图未必提升扰动预测准确性，却显著提升了时序调控重塑的可解释性——我们主张，这样一种诚实的边界重框，比又一个增量更准的预测器更能造福领域。")
P_rich("为从多角度探查所恢复的程序，我们三角验证三种独立的公开扰动模态（§3.7–3.8）中的方向性因果支持：谱系内批量 TF 敲除（L5a）、全基因组 L1000 过表达/敲低特征（L5b），以及脱离情境的癌细胞系中的单细胞分辨率 CRISPRi 筛选（L5c）。这一多模态设计，配上框定它的 K562 内部正对照，确立了程序级因果支持是“情境门控”的——这一发现重新框定了“对公开扰动数据的再分析”能够为所恢复 GRN 确证（以及不能确证）什么。")

# Figure 1
if os.path.exists(IMG):
    figp = doc.add_paragraph(); figp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figp.add_run().add_picture(IMG, width=Inches(5.8))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run("图 1. 面向恢复导向型虚拟扰动的、以证据驱动的框架。TSC-GNN 从状态条件化转录网络推断时序性重连的调控程序，并通过层级化证据框架加以解释，涵盖技术可重复性（L1）、生物学恢复（L2）、跨物种调控汇聚（L3）、转化假设生成（L4），以及情境门控的方向性因果支持（L5）。这些证据层共同逐步增强对所恢复调控程序的信心，同时明确界定框架的范围与局限。")
    set_run_font(r, size=9.5, italic=True)

# =====================================================================
# RESULTS
# =====================================================================
H("结果", 1)

H("3.1 预测准确性并不优于线性基线", 2)
P_rich("生物学发现建立在一个范围界定结果之上，它首先厘清该框架贡献什么、不贡献什么。我们在匹配输入下（相同的扰动向量 p、相同的训练划分、相同的读头）将图读头与线性基线进行基准比较，跨 **10 个随机种子 × 5 种图类型**（k-NN、DoRothEA、随机、置换 DoRothEA、0-hop）× 2 个任务 = 90 种配置**。在当前实现（固定因果图 + 线性读头）下，图组件在线性基线之外贡献了**有限的额外预测能力**：它在 **0 例**中胜过线性基线，在 **86 例**中无显著差异（图 2A）。一个非线性 KernelRidge 探针过拟合（相对误差 ≈ −126%），表明任何准确性增益都需要训练式的深度非线性读头，而非本文所研究的固定图。在外部 GEARS Perturb-seq 基准[9]上，图读头给出 −8.3%（单基因）/ −4.2%（双基因）的提升，二者置信区间均包含零。")
P_rich("由于图嵌入以线性方式进入读头，且与线性基线完全相同地提供，它在匹配输入下并未扩展线性假设空间；因此在该实现下，图并未带来更高的准确性。这一范围界定结果是下一步可解释性导向设计的前提。")
NOTE("【此处放置图 2A —— 预测基准：图 vs 线性，90 种配置】")

H("3.2 为何只有图才能恢复边级重连", 2)
P_rich("范围界定基准（§3.1）已确立：在线性读头下，固定因果图并不能提升*预测准确性*。这立即引出一个问题：**如果图不能提升预测，为何还要用图？** 因此，我们将这一贡献围绕*恢复*时序调控重连——即边级变化的恢复——而非围绕工具名称（TSC-GNN）来组织：框架是让恢复变得可解释的仪器，而非信息本身。")
P_rich("答案在于被测度的对象。一个从完整表达空间预测扰动结果的线性模型，将整个响应压缩为每个基因一个准确性数字。它无法将该预测变化归因于调控图中的*特定边*，因为它没有图：其输出是每个基因的标量 Δ，而非边级耦合变化。要恢复边级重连——*在某一生物转换中，哪个 TF→靶耦合增强或减弱*——必须将共表达变化分解到图基板上。这正是 ΔW（§5.3）所做的：它测度限制在每个有向边上的*Pearson 耦合变化*，这从根本上是一种图级操作。没有任何在原始表达上运行的线性模型能产生这一输出，因为耦合是定义在*边上*，而非*转录组上*。")
P_rich("TSC-GNN 最好被理解为不是一个**预测**引擎，而是一个**恢复**引擎：其交付物是带置换显著性的边级 ΔW 与模块级主调控因子排序——这些是线性预测器无法计算、基础模型读头无法分解的量。图之所以必要，不是因为它能做出更好的预测，而是因为它定义了*什么是可解释的*。其余结果展示了这一恢复视角所揭示的内容。")
NOTE("【此处放置图 2B —— 示意图：可解释性与预测的权衡】")

H("3.3 时序重连恢复了生物学上连贯、已确证的卒中程序", 2)
P_rich("跨四个转换应用状态条件化重连，恢复出一个连贯、与文献一致的卒中修复程序。在汇总 q < 0.05 下，sham→24 h、24 h→2 d、2 d→14 d、sham→14 d 分别有 **8、28、19、36 条边**显著（图 3A–D）。")
NOTE("【此处放置表 2 —— 各转换显著重连边汇总】")
B("急性损伤 onset（sham→24 h，8 条显著边）：损伤/炎症基因（如 *Tbx21→Cxcr3*，ΔW = −0.21）的耦合微弱但方向一致，与已知的急性缺血反应及轻微的急性信号相符。")
B("修复启动（24 h→2 d，28 条显著边）：最强的耦合增益是 *Sox2→Lsamp*（ΔW = +0.86，q < 0.001）与 *Sox10→Ank3*（ΔW = +0.74，q < 0.001）——少突胶质细胞谱系定向耦合的急剧上升，启动髓鞘再生。")
B("活跃髓鞘再生（2 d→14 d，19 条显著边）：典型少突胶质细胞髓鞘边 **Sox10→Plp1** 显著增强（ΔW = +0.51，q < 0.001），即在修复峰值处该方法独立恢复了髓鞘再生程序；*Hey2→Acta2*（ΔW = +0.59，q < 0.001）标志血管壁重塑。")
B("炎症消退（sham→14 d，36 条显著边）：*Sox9→Hapln1*（ΔW = +0.79，q < 0.001）与 *Sox10→Plp1*（ΔW = +0.42，q < 0.001）与修复正向耦合，追踪急性炎症程序的消退。")
P_rich("**已恢复程序的恢复（Level 2）。** 所恢复的主调控因子与已确证的卒中后修复生物学一致[23, 24, 25]。**Sox10、Sox2、Sox9**——少突胶质细胞谱系定向与髓鞘再生的典型主调控因子——驱动了最强的重连，并与顶级髓鞘再生边 *Sox10→Plp1* 形成闭环。作为负对照，在严格 q < 0.1 阈值下，跨全部三次分析唯一被复现的 TF 是 **Fos**——一种即早/AP-1 应激反应基因，被任何损伤非特异地诱导[26]；我们将其作为应激伪迹对照而非生物学信号呈现。")
P_rich("**所恢复靶点的合理性。** 显著重连边的靶点富集于少突胶质细胞（OR 36–61）、神经元（OR 110）与小胶质细胞（OR 15）标志物；15 个已知卒中相关 TF 中 14 个出现，且 77–83% 的边在校正 PC 组成后保留了方向——证实重连信号并非组成伪迹。PC 校正将原始→PC 方向翻转（如 *Sox10→Ank3*：原始 −0.38 → +0.73）表明，组成遮蔽可能隐藏真实的调控增益，而校正后的估计是更可靠的重度测量。")
NOTE("【此处放置图 3 —— 各转换的时序重连热图】")

H("3.4 第一级（L1）—— 跨队列主调控因子可重复性", 2)
P_rich("为检验技术稳定性，我们在两个独立队列上重新推导重连，并将 **TF 主调控因子排序** 与整合分析进行比较。由于整合的“sham”基线混合了两个队列的细胞，没有共享的转换且基因空间不同，我们在**主调控因子层面**评估可重复性（生物学信号在此稳健），而非在边层面（单条边估计本身噪声大）。整合排序汇集了全部四个转换，其中两个是跨队列拼接的伪转换（§4.3）；因此 ρ 反映的是通常驱动重连的 TF，其中 *Sox10* 是干净的跨队列复现范例。")
P_rich("主调控因子排序在队列间显著复现（表 3）：")
NOTE("【此处放置表 3 —— TF 排序的跨队列 Spearman 相关】")
add_table(
    ["比较", "共同 TF 数", "Spearman ρ", "p"],
    [
        ["整合分析 vs 队列 1", "251", "+0.517", "1.5 × 10⁻¹⁸"],
        ["整合分析 vs 队列 2", "235", "+0.548", "8.8 × 10⁻²⁰"],
        ["队列 1 vs 队列 2", "242", "+0.482", "1.7 × 10⁻¹⁵"],
    ],
    widths=[2.6, 1.6, 1.6, 1.8], font_size=9.5
)
P_rich("**Sox10** 出现在全部三次分析的 top-20 |ΔW|max TF 中；**Sox2/Sox9** 在 ≥2 次中出现。相比之下，单条边级显著性**未**跨队列复现（方向一致率 ≈ 0.52，即在随机水平；Jaccard ≈ 0），这与固定图/线性读头方法下单边估计固有的高方差一致，并作为预期行为在此报告，用以支撑模块级可解释性焦点，而非作为缺陷。")

H("3.5 第三级（L3）—— 所恢复程序的跨物种调控汇聚", 2)
NOTE("**范围披露：** 正交投影检验（§3.5a）考察小鼠恢复的 TF 的调控程序是否在人类 GRN 的*结构*中保守（通过独立的人类 DoRothEA 网络）。激活性检验（§3.5b）则分别考察相应靶程序是否在独立的人类卒中数据中*被转录激活*。二者共同将 L3 从“结构保守”提升为**跨物种调控汇聚**——在小鼠卒中中重连的同一程序（i）在人类 GRN 架构中保守，且（ii）在独立人类患者样本中共同激活。这是*功能汇聚*，而非*生物学验证*：批量队列来自外周血（非脑）且为横断面（非时间分辨），因此脑内固有和时间性的主张仍依赖于小鼠证据。")

H("3.5a 人类 GRN 中的结构汇聚", 3)
P_rich("我们取 12-TF 跨队列核心模块（在 3 次分析中复现 ≥2 次；**Sox10** 为三者共有），并将每个 TF 正交投影到其人类 DoRothEA 靶集上。对每个 TF，我们通过超几何检验测试两个文献整理参考集——*髓鞘/少突胶质细胞*[24, 27] 与 *神经炎症*[28, 29]——的富集，以 **2,000 次大小匹配（0.5–2× 靶计数）的随机人类 TF 置换**作为零分布。")
P_rich("三个链接显著保守（表 4）：")
NOTE("【此处放置表 4 —— 跨物种调控汇聚：结构富集】")
add_table(
    ["TF（人）", "参考基因集", "k/n", "OR", "超几何检验 p", "经验 p（置换）"],
    [
        ["SOX10", "髓鞘/少突", "7/22", "27.0", "6.0 × 10⁻⁸", "0.0005"],
        ["CEBPB", "神经炎症", "8/23", "16.5", "3.2 × 10⁻⁷", "0.024"],
        ["GATA2", "神经炎症", "15/23", "4.6", "3.3 × 10⁻⁴", "0.047"],
    ],
    widths=[1.1, 1.4, 0.8, 0.9, 1.6, 1.6], font_size=9
)
P_rich("**SOX10 → 人类髓鞘/少突：** 重叠 *PLP1、MBP、MAG、MPZ、PMP22*（主要髓鞘结构蛋白）外加 *GJC2*（少突胶质细胞间隙连接）与 *PDGFRA*（少突胶质细胞前体标志物）；小鼠顶级重连 TF 精确投射到人类髓鞘再生程序，与 Sox10→Plp1（§3.3）形成闭环。")
P_rich("**CEBPB → 人类神经炎症：** 重叠 *IL1B、IL6、TNF、CCL3、CCL5、NOS2、PTGS2、STAT3*。**GATA2 → 神经炎症：** 重叠 *TLR2、TLR4、NFKB1、NFKBIA* 以及 CCL/CXCL 趋化因子：即 TLR/NF-κB 天然免疫模块。")
P_rich("非显著/预期为阴性的结果被诚实报告：AR/ERG/NR2F2/PAX5/RUNX3/SOX9 未显示组织特异性富集（经验 p = 0.29–0.65）；SOX2/E2F1/GATA3 靶向过广，无法显示单组织富集（生物学上预期）。")

H("3.5b 在独立人类卒中批量数据（GSE16561）中的表达激活", 3)
P_rich("为超越结构保守，我们考察所恢复程序是否在独立人类卒中数据中被*转录激活*。我们获取了公开的白细胞全血批量 RNA-seq（GSE16561 [30]；39 例卒中 vs 24 例对照，Illumina HumanWG-6，RMA 标准化）。探针经 GPL6883 映射到 HGNC 符号，并归并为 17,493 个基因。对每个恢复的 TF，我们以其人类 DoRothEA 靶集，用 AUCell/ssGSEA 风格的秩–质量分数[6, 31] 为每个样本打分模块活性，再比较卒中 vs 对照（Mann–Whitney U；Cliff's δ；BH 校正[32]）。")
P_rich("三个程序均在人类卒中血液中显著激活（BH-q = 1.1 × 10⁻⁸），且强于 99.8% 的大小匹配随机基因集（经验 p = 0.002；表 5）：")
NOTE("【此处放置表 5 —— 人类卒中血液中所恢复程序的激活】")
add_table(
    ["模块", "靶基因（存在/总数）", "Δ（卒中−对照）", "Cliff's δ", "p"],
    [
        ["CEBPB（炎症）", "555/589", "+0.058", "+0.89", "3.8 × 10⁻⁹"],
        ["SOX10（髓鞘/少突）", "296/322", "+0.045", "+0.83", "4.9 × 10⁻⁸"],
        ["GATA2（炎症）", "4780/5370", "+0.023", "+0.76", "5.2 × 10⁻⁷"],
    ],
    widths=[1.8, 2.2, 1.6, 1.1, 1.5], font_size=9
)
P_rich("一个脑富集的负对照 TF（**PAX6**，344 个靶）*也*被激活（p = 4.5 × 10⁻⁸）。因此，所观察到的激活不能被解释为这三个调控程序在卒中中被*唯一*激活的证据；相反，这些发现表明部分信号可能反映了与系统性炎症及卒中后白细胞组成改变相关的广义转录重编程。因此，我们将 L3 结果解释为**跨物种调控汇聚的正交证据**：所恢复程序是人类卒中中被共同激活的模块之一，且其*方向*（炎症轴上调）与小鼠重连（§3.3）一致。外周血中 SOX10 模块的激活**不应**被解释为少突胶质细胞髓鞘再生的直接证据；SOX10 的少突胶质细胞特异性解读依赖于小鼠结构证据（§3.5a）。我们未做细胞组成反卷积，且基于血液的横断面设计限制了对脑内固有时间性重塑的直接推断。")
NOTE("【此处放置图 4 —— 跨物种调控汇聚：结构投影 + 血液激活】")

H("3.6 第四级（L4）—— 通过 LINCS 的药物扰动相关性", 2)
P_rich("作为转化演示，我们将 24 h 卒中损伤程序转换为疾病特征——上调的损伤/炎症基因（SPP1、CCL4、LGALS3、CD14、C5AR1、TNF、…）与一组异质性的低表达基因（干扰素反应、小胶质细胞标志物与转运基因）。值得注意的是，24 h 特征由**急性炎症/损伤轴**主导，本身并不包含下调的髓鞘/修复结构基因；髓鞘再生轴由重连分析独立恢复（Level 2–3，§3.3–3.5），而非由该特征得出。该特征源自 DoRothEA 网络基因空间（与重连相同的空间），即经调控程序过滤的差异表达，而非全转录组。")
P_rich("我们构建了两个特征：单队列**主**特征（up100/dn100，最佳分数 0.0671 且有严重并列，仅命中 1 个已知药物）与跨两队列相交的**稳健**一致特征（up28/dn17，最佳 0.125）。我们以**稳健**特征为主要结果；主特征作为单队列噪声对照，呼应 L1 可重复性主题。")
P_rich("稳健特征通过 L1000CDS2 反向匹配 API[22] 在 top 50 中返回 **4 个有文献支持的药物**：**vorinostat** 与 **trichostatin A**（HDAC 抑制剂），以及 **mevastatin** 与 **rosuvastatin**（他汀类）。它们的**抗炎**相关性与 24 h 特征所捕获的炎症轴直接吻合；它们的**促髓鞘/神经保护**相关性与重连分析（Sox10→Plp1，Level 2–3）独立恢复的髓鞘再生程序吻合——这是一种跨级汇聚，而非特征内效应。")
P_rich("**统计严谨性（置换背景）。** 使用唯一药物计数（避免因 L1000 多细胞系重复条目而虚增计数），20 个匹配大小的随机特征返回均值 **2.75 ± 1.37** 次命中（最大 5）。观测到的 **4** 次命中**未**超过随机（**经验 p ≈ 0.3**；N = 20，单次，实时 API）。一个初步的重复计数方案曾虚假提示 p = 0.048；我们纠正了这一点，并**不**主张显著富集。此外，vorinostat 的跨细胞系一致性（n_hits = 6）被证明是 L1000 背景效应：在随机特征下 vorinostat 出现在 53% 的运行中（最大 n_hits = 16）。")
NOTE("**[作者注 —— L4 局限]：** 第四级是一个**转化概念验证/假设生成**演示，*而非*药物疗效预测。LINCS 谱来自非神经元癌细胞系；该特征为单时间点（24 h 急性）且源自 DoRothEA 基因空间而非全转录组；L1000CDS2 显示高度并列；置换未达显著；且整个流程为纯计算机，无湿实验确认。")
NOTE("【此处放置图 5 —— 药物逆转：top 命中与置换分布】")

H("3.7 第五级（L5）—— 来自公开 TF 扰动的程序级方向性因果支持", 2)
NOTE("**范围披露：** L5 通过再分析**公开的、独立的** TF 功能缺失 RNA-seq，处理任何因果图方法的核心注意事项——*图从未被扰动*。我们在**程序（TF→靶模块）层面**检验：当 TF 自身被敲除时，框架所恢复的靶程序是否在下调基因中富集。这是**方向性因果支持**，明确**不是**边级验证（边不跨队列复现，§3.4），也**不是**关于卒中内因果性的主张（扰动数据非卒中）。")
P_rich("靶程序取自框架所用的**同一小鼠 DoRothEA GRN**（§5.2），因此独立于扰动数据集，检验非循环。")

H("3.7a 谱系内批量 KO（L5a）", 3)
P_rich("**Sox10 —— 少突胶质细胞特异性条件敲除（GSE269122 [33]）。** 所恢复的 Sox10 靶程序在被检候选中下调富集最强（表 6）：")
NOTE("【此处放置表 6 —— Sox10 cKO：靶程序方向性支持】")
add_table(
    ["TF 程序", "n", "平均 log₂FC", "OR↓", "Fisher p", "排名（最负）"],
    [
        ["Sox10（被扰动）", "319", "−0.090", "1.81", "≈ 0", "46 / 412（前 11%）"],
        ["Cebpb（对照）", "570", "−0.065", "1.83", "≈ 0", "147 / 412"],
        ["Gata2（对照）", "5177", "−0.033", "1.33", "≈ 0", "309 / 412"],
        ["Sox2（对照）", "3121", "−0.025", "1.29", "≈ 0", "339 / 412"],
    ],
    widths=[1.5, 0.8, 1.4, 0.9, 1.0, 1.8], font_size=9
)
P_rich("去除 Sox10 比所有其他候选程序更强地下调我们的 GRN 归因于 Sox10 的基因（排名 46/412；特异性梯度 Sox10 ≪ Cebpb < Gata2 < Sox2），与“去除一个激活因子 → 其靶基因下降”一致，并为 §3.3 恢复的 Sox10→髓鞘/髓鞘再生程序提供方向性因果支持。")
P_rich("**Cebpb —— 杂合敲除库普弗细胞（GSE273163 [34]）。** 所恢复的 Cebpb 程序再次在候选中下调富集最强，且对非扰动的 Sox10 程序具有干净特异性（表 7）：")
NOTE("【此处放置表 7 —— Cebpb het-KO：靶程序方向性支持】")
add_table(
    ["TF 程序", "n", "平均 log₂FC", "OR↓", "Fisher p", "排名（最负）"],
    [
        ["Cebpb（被扰动）", "555", "−0.172", "1.20", "0.022", "149 / 404（前 37%）"],
        ["Gata2（对照）", "4889", "−0.147", "1.15", "≈ 0", "205 / 404"],
        ["Sox10（对照）", "307", "−0.128", "0.95", "0.68", "278 / 404"],
    ],
    widths=[1.5, 0.8, 1.4, 0.9, 1.0, 1.8], font_size=9
)
P_rich("Cebpb 自身的靶基因显著下调（OR = 1.20，Fisher p = 0.022），且比非扰动的 Sox10 靶基因（OR = 0.95，p = 0.68）下调更多；效应温和是因为扰动为**杂合、3 vs 3** 设计，因此 L5 将其视为验证性方向支持而非决定性因果检验。")
P_rich("**跨数据集特异性。** Sox10 程序仅在 Sox10 被扰动处特异性下调（GSE269122 中排名 46 vs GSE273163 中 278）；Cebpb 程序在两处均下调（少突胶质细胞成熟与炎症程序共调控），但在 Cebpb 自身被扰动处下调最多。")

H("3.7b 独立基因扰动一致性 —— SigCom LINCS（L5b，A 模态）", 3)
P_rich("作为正交检验，我们查询 SigCom LINCS 基因扰动库[35]——CRISPR 敲低（l1000_xpr，140,603 个特征）与过表达（l1000_oe，33,782 个）——以每个 TF 的人类 DoRothEA 靶程序（最多 500 个靶，封顶）。对每个 TF，我们考察其*自身*扰动特征是否作为顶级逆转者（CRISPR KO）或顶级模拟者（OE）排名，以及该排名是否具有自我特异性。")
P_rich("最强的结果是 **GATA2 过表达**：GATA2 自身的 OE 特征作为 GATA2 靶程序的模拟者排名 **33,782 中第 3（前 0.01%）**（p = 1.4 × 10⁻⁵），且两个 GATA2 OE 重复均被分类为模拟者。自我特异性显著——自身百分位 0.01% vs 最佳跨 TF 4.98%（Δ = −4.97%）——提供有力正交证据：GATA2 是其 DoRothEA 恢复靶的激活因子。在 CRISPR-KO 库中，三个 TF 自身的 KO 特征均出现在逆转者的前 ~1%（SOX10 前 0.90%；CEBPB 前 0.46%；GATA2 前 1.17%），方向正确但**非自我特异**（自身 ≈ 跨 TF 百分位），可能因为 L1000 癌细胞系背景稀释了 TF 特异性效应。SOX10 OE 显示方向混合，且库中无 CEBPB OE 特征。")
NOTE("【此处放置表 8 —— SigCom LINCS：特征匹配结果】")

H("3.7c 单细胞分辨率 perturb-seq 再分析 —— Replogle 2022 K562（L5c，C 模态）", 3)
P_rich("作为更高分辨率、脱离情境的对照，我们再分析了 K562 中的全基因组 CRISPRi 筛选（Replogle et al., 2022 [36]；11,258 个扰动，585 个非靶向对照；Figshare 20029387），在该 TF 自身的 CRISPRi 特征（gemgroup Z 标准化伪批量）下检验每个 TF 恢复的 DoRothEA 靶程序。在这个脱离情境的癌细胞系中，检验返回**空结果**：三个 TF 的靶程序在自身扰动下均未被显著下移（表 9）。")
NOTE("【此处放置表 9 —— K562 sc-CRISPR：调控子响应结果】")
add_table(
    ["TF", "self-Z（靶向敲低）", "调控子平均 Z", "MWU p", "排名 / 332", "K562 背景"],
    [
        ["SOX10", "缺失（位点不在 K562）", "+0.018", "0.91", "304", "不表达"],
        ["CEBPB", "−0.43", "+0.007", "0.84", "239", "部分"],
        ["GATA2", "−0.19", "+0.004", "0.18", "139", "在背景内，但为空"],
    ],
    widths=[1.0, 1.9, 1.2, 0.9, 1.1, 1.2], font_size=9
)
P_rich("SOX10 的位点不在 K562 基因空间中（该 TF 在髓系 K562 中不表达）。CEBPB 与 GATA2 确认了靶向敲低（self-Z = −0.43 与 −0.19），但其靶程序的平均 Z 为 +0.007 与 +0.004（Mann–Whitney U p = 0.84 与 0.18），在全部候选程序中排名 239/332 与 139/332。")
P_rich("这一阴性结果具有信息性而非矛盾：它划定了程序级因果支持的**情境边界**。当扰动在生物学上恰当时——在原生谱系中敲除的 TF（L5a：Sox10 少突胶质细胞 cKO、Cebpb 库普弗细胞 KO）或以正确方向扰动的 TF（L5b：GATA2 过表达）——方向性支持被恢复，但在通用癌细胞系 CRISPRi 筛选中则不然，其中 TF 可能不活跃（SOX10）或其组织聚集的靶程序未被共调控（CEBPB/GATA2）。")

H("3.8 基因级扰动支持的跨模态综合（A + C）", 2)
P_rich("Level 5b 与 5c 用两种**独立的公开基因扰动模态**探查框架恢复的 TF→靶程序，二者在分辨率与*所测之物*上不同。联合分析产生任一单独模态都无法给出的洞见：**方向性因果支持是情境门控的，且两种模态将“特征模拟”与“调控子响应”分离开来。**")

H("3.8.1 方法", 3)
P_rich("**（M1）特征匹配模态 —— SigCom LINCS（A）。** 对每个焦点 TF，我们将其人类 DoRothEA 靶程序（上调基因查询，≤500 靶）提交至 SigCom LINCS 的 l1000_xpr 与 l1000_oe 库。我们记录该 TF*自身*扰动特征作为靶程序逆转者（KD）或模拟者（OE）的百分位排名，及其**自我 vs 跨 TF 特异性**。这测试转录组*尺度*的相似性：扰动 TF 是否产生全局上看起来像/反对靶模块的特征？")
P_rich("**（M2）调控子响应模态 —— Replogle 2022 K562 sc-CRISPR（C）。** 从全基因组 CRISPRi 伪批量（11,258 个扰动，585 个 NTC，gemgroup-Z；anndata 后端模式，Ensembl→符号经 mygene）中，我们取该 TF 自身的扰动行，以 NTC 均值做基线校正（np.nan_to_num 以抵消来自恒定方差基因的 gemgroup-Z 无穷），并检验**特定的 DoRothEA 靶集**是否被下移：靶平均 Z、单侧 Mann–Whitney U、底部四分位基因的 Fisher OR，以及在全部 332 个候选程序中的排名。这测试成员*解析*的响应：在内源敲低下，确切注释的靶是否移动？")
P_rich("**（M3）情境内正对照。** 为区分真正的细胞类型情境边界与失效的流程，我们在精选的 K562（红系–髓系白血病）主调控因子面板上运行 M2——GATA1、TAL1、KLF1、MYB、MYC、RUNX1、NFE2、BCL11A、ZBTB7A、FLI1、SPI1、CEBPA、E2F1——这些在自身 CRISPRi 下靶程序下移在生物学上预期。")
P_rich("**（M4）情境恰当性排序。** 我们将三个焦点 TF 的全部基因扰动证据按单一轴——扰动的生物学情境恰当性——排列，跨越**原生谱系**（L5a）→ **正确方向通用**（L5b）→ **脱离情境癌细胞系**（L5c），并读取支持的*梯度*。")

H("3.8.2 结果", 3)
P_rich("三个焦点 TF 的跨模态图谱。支持度随情境恰当性单调变化且依赖模态（表 10）。")
NOTE("【此处放置表 10 —— 跨模态扰动支持图谱】")
add_table(
    ["TF", "谱系内批量 KO（L5a）", "LINCS 过表达模拟（A）", "LINCS CRISPR-KD 逆转（A）", "K562 sc-CRISPR 调控子响应（C）"],
    [
        ["SOX10", "排名 46/412，OR↓ 1.81，p ≈ 0 ✓", "方向混合", "前 0.90%（非特异）", "排名 304/332 — 位点缺失"],
        ["CEBPB", "排名 149/404，OR↓ 1.20，p = 0.022 ✓", "无 OE 数据", "前 0.46%（非特异）", "排名 239/332，靶向，p = 0.84"],
        ["GATA2", "未检测", "排名 3/33,782，p = 1.4 × 10⁻⁵ ✓✓", "前 1.17%（非特异）", "排名 139/332，靶向，p = 0.18"],
    ],
    widths=[0.9, 1.9, 1.7, 1.6, 1.9], font_size=8.5
)
P_rich("**正对照确认流程在 K562 中有功效，但通用调控子是粗略代理（表 11）。** 在 K562 主 TF 中，调控子响应检验恢复了 **MYC（排名 19/332，MWU p = 3.1 × 10⁻³）** 与 **BCL11A（29/332，p = 7.5 × 10⁻³）** 的显著下移。三个焦点卒中 TF 在这些情境内阳性之下排名远低（139–304/332）。即便是典型的 K562 红系主因子 **GATA1**，也显示强靶向敲低（self-Z −0.54），却没有连贯的 DoRothEA 调控子下移（平均 Z +0.27，排名 187/332），表明即便在正确细胞类型中，通用 DoRothEA 调控子也只是部分追踪 TF 活性。")
NOTE("【此处放置表 11 —— K562 正对照调控子响应】")
P_rich("**结合 A + C 的新发现：**")
NUM("**两种独立模态的汇聚式情境门控。** 两个正交公开资源——LINCS 批量 L1000 与 Replogle K562 单细胞 CRISPRi——独立地未能为卒中程序提供 TF 特异性、脑相关的方向性支持，而谱系内 KO（L5a）则能。这些汇聚的脱离情境空结果作为**负对照**，排除了“靶程序可从*任何*扰动数据集恢复”这一平凡解释。")
NUM("**两种模态将特征模拟与调控子响应分离 —— 仅当结合时才可见。** GATA2 的过表达特征是第 3 强的模拟者（LINCS，前 0.01%），但其确切靶集在 K562 内源 CRISPRi 下并*不*移动（排名 139/332，p = 0.18）——即便 GATA2 是存在于 K562 的真正造血 TF。因此 LINCS 报告的转录组尺度“模拟/逆转”并*不是*该特定 DoRothEA 边在该背景中活跃的证据；只有成员解析的 Replogle 检验能裁断这一点，而它在脱离情境时为空。任一模态单独都无法揭示这一分离。")
NUM("**通过正对照校准解释：存在功效，但模块级检验有上限。** K562 内部阳性（MYC、BCL11A）证明空结果并非流程伪迹；GATA1 失败证明即便在情境内，通用调控子也只是粗略的、依赖广度与置信度的代理。这**限定**了任何模块级基因扰动检验——A 或 C——可被解释的强度。")
NUM("**梯度即结果，且正是可解释性优先框架所预测的。** 按情境恰当性排序，支持度单调衰减（原生谱系阳性 → 正确方向 OE 自我特异 → 脱离情境 CRISPRi 空），恰好符合“图编码的是*情境特异性*调控假设而非可迁移预测器”的预期。")
NOTE("【此处放置图 6 —— L5 三层三角验证：A + C 跨模态综合】")

# =====================================================================
# DISCUSSION
# =====================================================================
H("讨论", 1)

H("4.1 一个方法学边界，而非一个工具", 2)
P_rich("在固定因果图与线性解码器下，图结构在线性模型之外贡献**极少预测能力**。范围界定基准（§3.1）——90 种配置中图在零例胜过线性基线，且在外部 GEARS Perturb-seq 集上置信区间跨零——将其锚定在*固定图/线性读头*体制，而非工程失败。我们公开陈述而非掩饰它，因为该结论重新框定了图模型何时有用。")
P_rich("图的独特价值在于**可解释性**。有向因果图在单个 TF→靶边的分辨率上揭示调控*在何处*重塑，以及*哪些*主调控因子驱动该重塑。黑箱线性或基础模型读头将扰动压缩为单一分数并丢弃此结构。因此，我们将 TSC-GNN 定位为可解释性优先的虚拟扰动框架。（在 DoRothEA 与无向 k-NN 共表达图之间选择，正是基于这一理由：二者在预测上近乎无差异，但只有有向因果图能将重连呈现为“谁的调控被增强/减弱”这一可解释陈述。）")

H("4.2 预测为何失败 —— 以及为何这有意义", 2)
P_rich("GEARS 基准（§3.1）与 90 配置范围界定都指向同一结论：**在固定因果图 + 线性读头下，预测是一个线性问题。** 图并未扩展读头的假设空间，因为读头本身是线性的，且图作为固定的、预先计算的嵌入进入它——对图与线性基线完全相同。因此，该框架**不**声称比任何使用相同输入特征的其它方法更准确地预测扰动。")
P_rich("这一失败有意义，因为它揭示了根本权衡。**预测与解释是不同的任务，施加不同的架构需求：** 准确的扰动结果预测需要可学习的、非线性的、且可能无图的映射（基础模型[10, 11]、在学习图上训练的深度读头[37, 38]），而边级解释需要固定的、有向的因果图，将扰动效应分解到单个 TF→靶边。在本研究约束下，没有单一架构能同时最优服务两个目标。GEARS 结果——固定图相对无图基线表现为 −8.3%（单）/ −4.2%（双），置信区间含零——因此对该框架并非负面结果；它是将预测与解释分离、并各归其适合架构的诊断。这一不对称性——预测是线性的，解释需要图——在我们看来是本文最具普适性的方法学洞见，并直接关乎未来扰动建模系统如何设计。")

H("4.3 边不稳定性作为生物学发现，而非局限", 2)
P_rich("不稳定边与稳定程序之间的不对称性是本研究的核心生物学发现，它重新框定了图买什么、不买什么。五个验证级映射到五个主张：")
NUM("**预测并非图的优势。** 在匹配输入下，固定因果图在线性基线之外贡献有限预测能力（§3.1）。")
NUM("**可解释性是图的优势。** 图结构独特地启用带置换显著性的边级重连 ΔW 与模块级主调控因子推断（§3.3–3.4）。")
NUM("**恢复的程序复现已知生物学。** 重连模块恢复了时序有序的修复程序——炎症消退、少突胶质细胞髓鞘再生、神经元恢复——且核心 TF（Sox10/Sox2/Sox9）是已确证的修复调控因子（§3.3，L2）。")
NUM("**程序是保守的、可转化的、且方向性因果支持的。** 调控架构投射到人类保守的卒中修复程序（SOX10→髓鞘，OR = 27；CEBPB/GATA2→神经炎症；§3.5a，L3），并在公开人类卒中血液中*独立*激活（§3.5b）——跨物种调控汇聚的正交证据。对公开 TF 功能缺失 RNA-seq 的再分析（§3.7）显示程序级的方向性因果支持。跨模态综合（§3.8）增加了两个关键校准：（i）特征模拟（LINCS）与调控子响应（K562 CRISPRi）*分离*，因此全局模拟信号本身并不确证特定边活跃；（ii）K562 内部正对照确认脱离情境空结果是真正的细胞类型边界，同时也暴露上限——通用调控子即便在情境内也只是粗略代理（GATA1 失败）。这些层共同从可解释虚拟扰动闭环到因果支持、可转化的假设。")
P_rich("**边不稳定性作为生物学发现，而非局限。** 框架被描述为交付边级重连 ΔW（§3.3），然而全部五个验证级（L1–L5）都在*模块*（程序）层面运作。这一并置——边级输出、模块级验证——是刻意的，并构成一项关键发现本身。单个边估计（如 Sox10→Ank3，ΔW = +0.74）足够精确以引导生物学注意并排序主调控因子，但它们**不**跨队列复现（单条边 Jaccard ≈ 0；方向一致率 ≈ 0.52，在随机水平；§3.4）。相比之下，L3–L5 的模块级证据显示，*聚合的* TF→靶程序跨队列保守、跨物种保守，并受独立扰动数据方向性支持。这一不对称性——**边是数据集特异的；程序是稳健的**——具有生物学意义：它暗示单个 TF→靶耦合是针对细胞背景与实验平台精细调谐的，而*调控逻辑*（一个 TF 控制哪些靶）是硬编码的。该发现映照网络生物学中公认的原理：生物网络中的单边固有噪声且依赖背景，但网络级属性（模块成员、主调控因子排名）稳定[18, 48, 69]。因此，我们公开报告 Jaccard ≈ 0 结果，不是作为要隐藏的缺陷，而是作为这一洞见的经验基础——并作为对领域的警示：任何 GRN 方法中的单边主张都应被解释为数据集特异的假设，而非普适发现。这一区分映照领域与 GRN 方法的历史经验[20, 48]：声称边级因果性的方法（ARACNE，2006 [70]）未能维持采用，而提供模块级假设生成的方法（SCENIC [6]、NicheNet [8]）被广泛使用。TSC-GNN 遵循后一范式，并诚实披露其边级输出。")
NUM("**边界即信息，而非弱点。** 固定因果图在匹配输入下不提升预测的发现不应被隐藏。在我们看来，这是本研究最具普适性的洞见：领域朝向更大、更不透明的扰动模型的轨迹，可能绕开边级重连所提供的机制层面理解。")
P_rich("**为何程序存续。** 边级不稳定下程序级重连的稳定性（§3.4）由生物基因调控网络的三个属性解释。第一，**网络简并性**[71, 72]：多个不同的边配置可收敛于同一调控结果，因此像 Sox10 这样的 TF 并非通过单一不变边控制其髓鞘靶，而是通过分布式、冗余的相互作用集合，其中任一子集可在给定背景中活跃。第二，共调控因子间的**功能冗余**[73]：Sox10、Sox2、Sox9 与 Olig2 共享其髓鞘靶的相当部分，因此当某一队列中 Sox10→X 边弱时，Sox2→X 或 Sox9→X 补偿，程序级信号得以保留。第三，**吸引子动力学**[74, 75]：GRN 拥有稳定的吸引子状态，不同网络配置都向其松弛，因此髓鞘/少突胶质细胞程序（Sox10 靶集）在不同背景中通过不同接线重构——类似于在噪声单基因表达下细胞类型身份稳健性的著名现象。因此，所恢复程序由网络架构缓冲，而非钉在单个边上。这一简并性–稳定性不对称性本身是该框架的关键生物学意涵，也是纯预测模型无法呈现的：仅报告单一准确性数字的模型丢弃了使该不对称性可见的边分辨率结构。正是在这里——揭示机制而非最大化预测——TSC-GNN 的价值所在。")
P_rich("**预期审稿人问题。** 我们回应审查中可能出现的三个问题。")
P_rich("*Q1：若 GNN 不优于线性基线，为何还要用 GNN？* 范围界定基准（§3.1）与 §4.2 直接回答：GNN 贡献的是解释能力，而非预测准确性。固定图 GNN + 线性读头，与无图线性基线，运行于同一预测假设空间；它们仅在*何处*与*哪些*调控变化被揭示上不同。图之必要不在于预测，而在于将扰动效应分解到单个边——这是任何无显式图的线性或基础模型读头都无法执行的任务。该框架被设计为恢复引擎，而非预测引擎。")
P_rich("*Q2：若跨队列边级重叠近零（Jaccard ≈ 0），为何应信任边级重连？* 边级估计是数据集特异的假设，而非已验证的发现。模块级证据（L3–L5）确立*聚合程序*方向正确且跨队列、跨物种、跨扰动模态保守。单个边可解释为生成可检验假设的、依赖背景的耦合变化，但不应被视为不变的生物学真理。这一区分——边有噪声、程序稳健——本身是框架的关键生物学发现（§4.3，“边不稳定性作为生物学发现”）。")
P_rich("*Q3：该框架是否可推广到卒中之外？* 该框架与疾病无关，可应用于任何具有 DoRothEA 兼容调控子库的多时间点单细胞数据集。可能有成效的应用包括脊髓损伤（SCI）、多发性硬化（MS）、创伤性脑损伤（TBI）与慢性炎症——它们都涉及可通过相同计算组件（时序 GRN、伪时间、CellChat）触及的胶质与免疫 GRN 的时序重塑。拼接时间轴的需求与固定图局限将同样适用；L5 因果支持范式可采纳于任何存在相关细胞类型公开 TF 扰动数据的疾病。")

H("4.4 局限", 2)
P_rich("若干约束是当前设计固有的。")
P_rich("**时间轴。** 纵向轴拼接自两个队列（GSE174574 与 GSE225948；§5.1）。所得轴非均匀采样（缺失 3–7 d 窗口）且带有批次×时间混淆。拼接段间的离散 ΔW 并非连续速率，且跨队列边级比较受不同基因空间限制（6,897 / 8,736 / 7,041 个基因）。")
P_rich("**L4 药物逆转。** 药物逆转级是假设生成演示，而非疗效预测。LINCS 谱来自非神经元癌细胞系；置换未达显著（p ≈ 0.3）；且整个流程为纯计算机，无湿实验确认（§3.6）。")
P_rich("**L5 因果支持。** L5 级是对*公开、非卒中扰动数据在程序层面的再分析*。它支持所恢复 TF→靶程序的方向性，但不是卒中内的因果验证、不是边级主张、也不是对框架特定边的直接检验（§3.7）。SigCom LINCS 扩展为 GATA2 提供正交方向性证据（OE 排名 3/33,782，强自我特异性），并为三个 TF 的 CRISPR KO 提供正确的逆转方向，但在 CRISPR-KO 库中缺乏 TF 级自我特异性，且 SOX10 OE 显示方向混合。K562 sc-CRISPRI 扩展在该脱离情境系中返回空——SOX10 在髓系 K562 中不表达，且 CEBPB/GATA2 被靶向敲低但其靶程序未被共下调——我们透明报告为*情境边界*而非反驳。K562 内部正对照支持这一解读，但也暴露上限：即便典型 K562 红系主因子 GATA1 也显示靶向敲低而无连贯 DoRothEA 调控子响应，因此通用调控子只是粗略代理，限定了任何模块级基因扰动检验可被解释的强度（§3.8）。")
P_rich("**L3 物种汇聚。** 人类激活检验使用外周血（非脑）、为横断面（非时间分辨），且未做细胞组成反卷积。血液中 SOX10 模块激活不能被解释为少突胶质细胞髓鞘再生（§3.5b）。PAX6（脑特异性负对照）也被激活，表明部分信号反映广义转录重编程。")
P_rich("**模型能力与可重复性基础设施。** 全部分析使用固定因果图（DoRothEA）+ 线性读头。可学习图或非线性深度读头可能改变预测与可解释性结论。我们刻意将自己约束于固定图/线性读头体制以隔离图结构的贡献，但未来工作应扩展到学习图与深度读头。在可重复性方面，当前流程使用 `mygene` [https://mygene.info] 查询做 Ensembl→符号映射，首次运行需联网；我们提供预计算缓存作为补充表 S4，使流程在初始映射后可完全离线执行。`bbb_gnn` conda 环境锁定全部依赖（numpy 2.2.6、scipy 1.15.3、pandas 2.3.3），分析清单包含缓存状态的 SHA-256 哈希；我们建议每次可重复性运行前新建 `conda env create -f environment.yml`。领域与 GRN 方法的历史[20, 48]表明，采用度与基础设施（pip 可安装包、Docker 镜像、教程 notebook）同样取决于科学性能：我们以 GitHub 仓库形式提供流程，配 Binder 就绪 notebook，并计划推出 conda-forge 包。")
P_rich("**外部效度。** 该框架在单一疾病背景（缺血性卒中）中演示。向其它时序重连领域（发育、神经退行、癌症）的普适性仍有待检验。")

H("4.5 未来方向", 2)
P_rich("未来工作可沿三个轴扩展框架：")
NUM("**可学习图架构**，其中 GRN 部分或完全从数据学习，可能使用基于注意力的图学习[37]或可微分因果发现[38]。这可能恢复静态 DoRothEA 先验遗漏的细胞类型特异边，同时保持经由边级 ΔW 读头的可解释性。")
NUM("**非线性、训练式深度读头**（如图 Transformer 解码器），可利用图结构获得更高预测准确性，在固定图/线性读头体制表现不佳的外部 Perturb-seq 基准上检验。")
NUM("**应用到其它疾病与发育背景**（神经退行、脊髓损伤、癌症进展），以检验时序重连框架与 L5 因果支持范式的普适性。")

H("4.6 结论", 2)
P_rich("固定因果图经线性解码器读出，在预测上并不胜过线性模型；这一边界是本工作的起点而非失败之处。图*所*提供的是**可解释性**：带置换显著性的边级重连（ΔW）、模块级主调控因子排序、细胞间通讯重塑，以及药物重定位假设——从虚拟*计算机*扰动到因果支持、可转化输出的闭环。五级证据阶梯（L1–L5）分别确立跨队列可重复性、生物学连贯性、跨物种调控汇聚、转化汇聚，以及程序级方向性因果支持，同时坦率暴露每一级的边界。我们主张，诚实的否定预测结果与依赖背景的因果支持梯度，比又一个增量更准的预测器，是对领域更持久的贡献。")

# =====================================================================
# METHODS
# =====================================================================
H("方法", 1)

H("5.1 数据来源", 2)
P_rich("**小鼠缺血性卒中单细胞转录组。** 我们使用了两个独立生成的小鼠 MCAO scRNA-seq 队列。**队列 1 —— GSE174574**（Li et al., 2021 [3]）：3 个 MCAO（24 h）+ 3 个 sham，57,528 个细胞，含 BEAM 伪时间轨迹。**队列 2 —— GSE225948**（Anrather et al., 2024 [4]）：真正的双时间点卒中后时间进程（2 d 与 14 d，脑与血）。没有单一公开 scRNA-seq 队列覆盖全部三阶段；我们通过将队列 1（24 h + sham）与队列 2（2 d / 14 d）拼接，组装了 24 h → 2 d → 14 d 轴，跨越急性 → 亚急性峰值 → 修复/重塑阶段。这是一局限（§4.3）。")
P_rich("**人类卒中批量转录组。** GSE16561 [30]：39 例卒中 vs 24 例对照，外周全血，Illumina HumanWG-6。")
P_rich("**公开 TF 敲除数据。** GSE269122 [33]：少突胶质细胞中 Sox10 条件敲除（4 KO / 4 WT，胼胝体）。GSE273163 [34]：库普弗细胞中 Cebpb 杂合敲除（3 Hete / 3 WT）。")
P_rich("**单细胞 CRISPRi 筛选。** Replogle et al. 2022 [36]，K562 全基因组 CRISPRi，gemgroup Z 标准化伪批量（Figshare 20029387，文件 K562_gwps_normalized_bulk_01.h5ad；11,258 个扰动，585 个 NTC）。")
P_rich("**预处理。** 计数矩阵经文库大小归一化（×(10⁴)）后 log1p 加载并处理，在固定 conda 环境（bbb_gnn；numpy 2.2.6 / scipy 1.15.3 / pandas 2.3.3）中。每个细胞保留细胞类型组成以做组成校正。每个队列在相同协议下处理。")

H("5.2 基因调控网络构建", 2)
P_rich("我们使用 **DoRothEA** 共识调控子[20]（置信水平 A–C）作为*有向因果* TF→靶图——方向与符号（激活/抑制）取自文献整理先验，因此我们报告的任何重连都可解释为“谁的调控被增强/减弱”，而非仅仅是无向共表达变化。小鼠与人类调控子从本地 TSV 导出读取，使流程完全离线且可复现（每次运行含 SHA-256 清单）。对每个 TF，我们额外计算**状态亲和**向量 A_aff（n=4,000 细胞的子样本），捕捉该 TF 的靶在各细胞状态中表达强度；按 |A_aff| 排名前 50% 的边保留为*状态条件化*边用于重连检验。")

H("5.3 TSC-GNN：状态条件化边重连", 2)
P_rich("对每个被测有向边 e = (TF u → 靶 v) 与每个时间点 t，我们计算**状态内 Pearson 耦合**")
MATH("r_{e,t} = corr(x_u^{(t)}, x_v^{(t)})")
P_rich("其中 x_u^{(t)}, x_v^{(t)} 是限制于状态/时间 t 细胞的 u 与 v 的 log 归一化表达。转换 t_1 → t_2 的**重连效应**为耦合变化")
MATH("ΔW_{e, t_1→t_2} = r_{e,t_2} − r_{e,t_1}")
P_rich("正 ΔW 表示*耦合增强*；负表示*耦合减弱*。")
P_rich("**组成校正。** 我们对 x_u, x_v 关于全表达矩阵的前 n_pc = 10 个主成分（最小二乘回归）做残差化，使 ΔW 反映去除主组成变化轴后的耦合变化。原始与 PC 校正 ΔW 均报告。")
P_rich("**置换检验与多重检验。** 对每个转换，我们置换时间标签（n_perm = 200，seed = 2）并重新计算 ΔW 以构建每条边的零分布；每条边的双侧 p 值为 p_e = (1 + #{perm: |ΔW^{null}_e| ≥ |ΔW^{obs}_e|}) / (n_perm + 1)。我们选择 n_perm = 200 以在计算成本（单 CPU 每转换约 25 分钟）与可获得最小 p 值（0.005）间取得平衡，这对我们典型的边计数（每转换 ≤36 条显著边）已足够。对显著边多得多的数据集，我们建议增加 n_perm 或切换到 §5.9 描述的解析近似。我们应用两种校正：（i）每转换的 **Benjamini–Hochberg FDR**；（ii）**置换合并 FDR**，以每条边的 |ΔW| 除以其零分布标准差做标准化，跨边与置换汇集零 z 分数，并计算不依赖 p 分辨率的 q 值。合并 q < 0.05 的边报告为显著重连。")

H("5.4 细胞间通讯重塑", 2)
P_rich("细胞间通讯分析遵循 CellChat 框架[21]，分别应用于每个时间点，以识别跨越卒中转换强度变化的配体–受体对。")

H("5.5 第三级：跨物种检验", 2)
P_rich("(a) **结构富集。** 对每个小鼠恢复的 TF，我们将其 DoRothEA 靶集正交投影到人类同源基因（经 HGNC）。通过超几何检验，以 2,000 次大小匹配随机人类 TF 置换，检验文献整理的髓鞘/少突胶质细胞与神经炎症参考集的富集。")
P_rich("(b) **表达激活。** 模块活性经 AUCell/ssGSEA 秩–质量分数[6, 31]在 GSE16561 中每个 TF 的人类 DoRothEA 靶集上打分，比较卒中 vs 对照（Mann–Whitney U；Cliff's δ；BH 校正）。")

H("5.6 第四级：通过 LINCS 的药物逆转", 2)
P_rich("24 h 卒中损伤特征（上/下调基因集）提交至 L1000CDS2 API[22] 做反向连接映射。稳健特征（跨两队列相交，up28/dn17）为主要结果。置换背景由 20 个随机大小匹配特征（唯一药物计数）估计。")

H("5.7 第五级：公开 TF 扰动再分析", 2)
P_rich("(a) **谱系内批量 KO。** 对每个 TF，我们从公开 DE 表计算其 DoRothEA 靶程序的平均 log₂FC、底部四分位富集的 odds ratio 与 Fisher 精确 p，以及在全部 DoRothEA TF 程序中的排名。")
P_rich("(b) **SigCom LINCS（A 模态）。** 我们通过 SigCom LINCS 数据 API[35] 查询 l1000_xpr（CRISPR-KD，140,603 个特征）与 l1000_oe（过表达，33,782 个），提交每个 TF 的人类 DoRothEA 靶程序（最多 500 靶）。database 参数作为字符串库名传递。")
P_rich("(c) **Replogle K562 sc-CRISPR（C 模态）。** 从 h5ad（后端模式，anndata 0.11.4），我们通过解析 obs.index（<ID>_<GENE>_<guide>_<ENSG>）识别扰动行。NTC = non-targeting（n = 585）。var ID（Ensembl）经 mygene（querymany，species = human）映射到基因符号。对每个 TF，我们：（i）提取其自身扰动行；（ii）用 np.nan_to_num 计算 NTC 基线均值以抵消来自恒定方差基因的 gemgroup-Z 无穷；（iii）计算效应 = 扰动均值 − NTC 均值；（iv）用 Mann–Whitney U（单侧，alternative = “less”）、底部四分位 Fisher OR，以及在全部 332 个候选程序中的排名，检验靶基因下移。")

H("5.8 预测基准", 2)
P_rich("我们基准比较了 10 随机种子 × 5 图类型（k-NN、DoRothEA、随机、置换 DoRothEA、0-hop）× 2 任务 = 90 种配置。线性基线接收匹配输入（相同扰动向量 p、相同训练划分）。相对提升计算为（图 MSE − 线性 MSE）/ 线性 MSE。")

H("5.9 可重复性", 2)
P_rich("全部运行输出一份清单（命令、库版本、种子、缓存的 SHA-256）。在固定种子下分析完全确定。完整源代码与确切分析脚本发布于数据与代码可用性中描述的公开仓库。")

H("5.10 命名与架构理由", 2)
P_rich("该框架命名为 TSC-GNN（时序/细胞状态条件化图神经网络）。其范围值得直白陈述，因为它解决审稿人反复提出的问题。TSC-GNN 运行于**固定的、文献整理的因果图**（DoRothEA，§5.2），并通过**线性**解码器读出边级重连 ΔW——状态条件化的 Pearson 耦合变化带置换显著性（§5.3）。它**不**训练具有可学习节点或边嵌入的消息传递网络。这是刻意的设计选择，而非局限：§3.2 表明固定图 + 线性读头已提供我们寻求的可解释性（边级 ΔW），且 §3.1 证明在学习图在匹配输入下并无预测优势。消息传递 GNN 将引入不可解释的参数，而按我们的基准，它不带来准确性增益——直接用框架唯一贡献（可解释性）去换取并不存在的收益。因此，TSC-GNN 中的“图”指的是**结构化因果基板**，虚拟扰动在其上传播并被分解，而非端到端学习的 GNN。这将 TSC-GNN 定位为*恢复*引擎（参 §1、§4.1）而非*预测*引擎，并使其区别于端到端扰动 GNN 如 GEARS[9]。")

# =====================================================================
# DATA AND CODE AVAILABILITY
# =====================================================================
H("数据与代码可用性", 1)
P_rich("**数据。** 全部原始组学数据公开可用，且我们在不生成新测序的情况下进行了再分析：小鼠卒中 scRNA-seq，GSE174574 [3] 与 GSE225948 [4]；人类卒中血液批量，GSE16561 [30]；TF 敲除分析，GSE269122 [33]（Sox10）与 GSE273163 [34]（Cebpb）；K562 全基因组 CRISPRi，Figshare 20029387 [36]。")
P_rich("**代码。** 完整 TSC-GNN 流程——预处理、GRN 构建、带置换显著性的边级重连、CellChat 重塑、LINCS 药物逆转，以及第五级扰动再分析——以 Python 3 实现，依 MIT 许可发布。源代码与确切分析脚本发布于 https://github.com/ming13333/tsc-gnn（接受后发布）并归档于 Zenodo（DOI：10.5281/zenodo.XXXXXXX）。仓库附带 conda 环境文件（environment.yml）、每次运行的 SHA-256 清单，以及子采样队列上的工作示例。图 3–7 与表 1–11 背后的处理表作为补充数据提供。")
P_rich("**注。** GitHub 与 Zenodo 链接为占位符，将在接受后定稿；代码亦可应要求从通讯作者处获取。")

H("作者贡献", 1)
P_rich("Ming Luo 与 Yimin Mei 对本工作贡献相同。Wenwu Huang 为通讯作者（hwenwu321@gmail.com）。各作者在构思、方法开发、数据分析与论文撰写中的角色分配将由作者最终确定。")

H("致谢", 1)
P_rich("[待作者补充。]")

H("利益竞争", 1)
P_rich("作者声明无利益冲突。")

# =====================================================================
# REFERENCES (kept verbatim in English)
# =====================================================================
H("参考文献", 1)
NOTE("**[作者注：全部参考文献为草稿条目。作者须在投稿前对照发表版本逐条核实。]**")
refs = [
    "[1] Cramer, S. C. & Carrico, C. R. Stroke recovery: a perspective on mechanisms. Nat. Rev. Neurosci. 9, 720–731 (2008). — Stroke recovery mechanisms overview.",
    "[2] Carmichael, S. T. The 3 Rs of stroke biology: repair, regeneration, remodelling. Nat. Rev. Neurosci. 17, 420–432 (2016). — Temporal phases of stroke recovery.",
    "[3] Li, S. et al. Single-cell transcriptomic analysis of ischemic stroke reveals the temporal dynamics of glial and neuronal responses. Acta Neuropathol. Commun. 9, 152 (2021). — GSE174574.",
    "[4] Anrather, J. et al. Single-cell RNA sequencing of the post-stroke brain and blood reveals a temporally coordinated immune response. Nat. Immunol. 25, 294–307 (2024). — GSE225948.",
    "[5] Kamimoto, K. et al. Dissecting cell identity via network inference and in silico gene perturbation. Nature 614, 742–751 (2023). — CellOracle.",
    "[6] Aibar, S. et al. SCENIC: single-cell regulatory network inference and clustering. Nat. Methods 14, 1083–1086 (2017). — SCENIC / AUCell.",
    "[7] Bravo González-Blas, C. et al. SCENIC+: single-cell multiomic inference of enhancer-driven gene regulatory networks. Nat. Methods 20, 1355–1367 (2023). — SCENIC+.",
    "[8] Browaeys, R., Saelens, W. & Saeys, Y. NicheNet: modeling intercellular communication by linking ligands to target genes. Nat. Methods 17, 159–162 (2020). — NicheNet.",
    "[9] Roohani, Y., Huang, K. & Leskovec, J. Predicting transcriptional outcomes of novel multigene perturbations with GEARS. Nat. Biotechnol. 42, 157–166 (2023). — GEARS.",
    "[10] Cui, H. et al. scGPT: toward building a foundation model for single-cell multi-omics. Nat. Methods 21, 1469–1480 (2024). — scGPT.",
    "[11] Hao, M. et al. Large-scale foundation model for single-cell transcriptomics. Nat. Methods 21, 1481–1491 (2024). — scFoundation.",
    "[12] Wang, W., Hu, Z., Weiler, P., Mayes, S., Lange, M., Fountain, D. M., Haug, J. O., Wang, J., Xue, Z., Sauka-Spengler, T. & Theis, F. J. RegVelo: gene-regulatory-informed dynamics of single cells. Cell 189, 3773–3800.e44 (2026). — RegVelo.",
    "[13] Liu, Y., Cai, F., Barile, M., Chang, Y., Cao, D. & Huang, Y. TemporalVAE: atlas-assisted temporal mapping of time-series single-cell transcriptomes during embryogenesis. Nat. Cell Biol. 27, 1982–1992 (2025). — TemporalVAE.",
    "[14] Ishikawa, M. et al. RENGE infers gene regulatory networks using time-series single-cell RNA-seq data with CRISPR perturbations. Commun. Biol. 6, 1290 (2023).",
    "[15] Kipf, T. N. & Welling, M. Semi-supervised classification with graph convolutional networks. Proc. ICLR (2017). — GCN.",
    "[16] Veličković, P. et al. Graph attention networks. Proc. ICLR (2018). — GAT.",
    "[17] Hamilton, W. L., Ying, Z. & Leskovec, J. Inductive representation learning on large graphs. Proc. NeurIPS (2017). — GraphSAGE.",
    "[18] Dwivedi, V. P. et al. Graph neural networks: a review of methods and applications. NeurIPS (2023). — GNN review.",
    "[19] Lotfollahi, M. et al. Predicting cellular responses to perturbations with scGen. Nat. Methods 16, 1253–1261 (2019). — scGen.",
    "[20] Garcia-Alonso, L. et al. Benchmark and integration of resources for the estimation of human transcription factor activities. Genome Res. 29, 1363–1375 (2019). — DoRothEA.",
    "[21] Jin, S. et al. Inference and analysis of cell–cell communication using CellChat. Nat. Commun. 12, 1088 (2021). — CellChat.",
    "[22] Duan, Q. et al. L1000CDS2: LINCS L1000 characteristic direction signatures search engine. npj Syst. Biol. Appl. 2, 16015 (2016). — L1000CDS2.",
    "[23] Stolt, C. C. et al. Terminal differentiation of myelin-forming oligodendrocytes depends on the transcription factor Sox10. Genes Dev. 16, 165–170 (2002). — Sox10 role.",
    "[24] Nave, K. A. Myelination and support of axonal integrity by glia. Nat. Rev. Neurosci. 11, 275–283 (2010). — Myelin biology.",
    "[25] Fancy, S. P. J. et al. Myelin regeneration: a review of the cellular and molecular mechanisms. Ann. Neurol. 69, 579–589 (2011). — Remyelination.",
    "[26] Morgan, J. I. & Curran, T. Stimulus-transcription coupling in the nervous system: involvement of the inducible proto-oncogenes fos and jun. Annu. Rev. Neurosci. 14, 421–451 (1991). — Fos stress response.",
    "[27] Zawadzka, M. et al. CNS-resident glial progenitor/stem cells produce Schwann cells as well as oligodendrocytes during repair of CNS demyelination. Cell Stem Cell 6, 578–590 (2010). — Oligodendrocyte progenitor.",
    "[28] Doyle, K. P., Simon, R. P. & Stenzel-Poore, M. P. Mechanisms of ischemic brain damage and repair. Stroke 39, 571–578 (2008). — Neuroinflammation mechanisms.",
    "[29] Jin, R., Yang, G. & Li, G. Inflammatory mechanisms in ischemic stroke. Prog. Neurobiol. 90, 178–189 (2010). — Inflammation in stroke.",
    "[30] Barr, T. L. et al. Genomic expression in acute stroke patients. Stroke 41, 2280–2285 (2010). — GSE16561.",
    "[31] Barbie, D. A. et al. Systematic RNA interference reveals that oncogenic KRAS-driven cancers require TBK1. Nature 462, 108–112 (2009). — ssGSEA.",
    "[32] Benjamini, Y. & Hochberg, Y. Controlling the false discovery rate: a practical and powerful approach to multiple testing. J. R. Stat. Soc. B 57, 289–300 (1995). — BH correction.",
    "[33] Jörg, L. M. et al. Transcription factors Sox8 and Sox10 contribute with different importance to the maintenance of mature oligodendrocytes. Int. J. Mol. Sci. 25, 8754 (2024).",
    "[34] Lin, S.-Z. et al. C/EBPβ–VCAM1 axis in Kupffer cells promotes hepatic inflammation in MASLD. JHEP Rep. 7, 101418 (2025).",
    "[35] Subramanian, A. et al. A next generation connectivity map: L1000 platform and the first 1,000,000 profiles. Cell 171, 1437–1452.e17 (2017). — LINCS L1000 / SigCom.",
    "[36] Replogle, J. M. et al. Mapping information-rich genotype–phenotype landscapes with genome-scale Perturb-seq. Cell 185, 2559–2575.e28 (2022). — K562 genome-scale CRISPRi.",
    "[37] Veličković, P. Message passing all you need: graph attention networks and beyond. Nat. Rev. Phys. 5, 343–356 (2023). — Graph learning.",
    "[38] Zheng, X. et al. DAGs with NO TEARS: continuous optimization for structure learning. Proc. NeurIPS (2018). — Differentiable causal discovery.",
    "[39] Lotfollahi, M. et al. Biologically informed deep learning for perturbation prediction and biological discovery. Nat. Biotechnol. 41, 1759–1770 (2023). — scPerturb.",
    "[40] Bereket, M. & Karaletsos, T. PerturbNet: a graph neural network for predicting perturbation outcomes. [bioRxiv] (2022). — PerturbNet.",
    "[41] Peidli, I. et al. scPerturb: harmonized single-cell perturbation data. Nat. Biotechnol. 42, 1311–1319 (2024). — scPerturb compendium.",
    "[42] Singh, R. et al. PerturbExpress: a benchmarking platform for perturbation prediction. Nat. Microbiol. (2024).",
    "[43] Barry, C. et al. BEAM: branching expression analysis modeling for pseudotime inference. Nat. Biotechnol. 40, 409–415 (2022). — BEAM pseudotime.",
    "[44] Street, K. et al. Slingshot: cell lineage and pseudotime inference for single-cell transcriptomics. BMC Genomics 19, 477 (2018). — Slingshot.",
    "[45] Cao, J. et al. The single-cell transcriptional landscape of mammalian organogenesis. Nature 566, 496–502 (2019). — Monocle 3.",
    "[46] Efremova, M. et al. CellPhoneDB: inferring cell–cell communication from combined expression of multi-subunit ligand–receptor complexes. Nat. Protoc. 15, 1484–1506 (2020).",
    "[47] Badia-i-Mompel, P. et al. decoupler: ensemble of computational methods to infer biological activities from omics data. Bioinformatics 38, 3631–3634 (2022). — decoupler.",
    "[48] Pratapa, A. et al. Benchmarking algorithms for gene regulatory network inference from single-cell transcriptomic data. Nat. Methods 17, 147–154 (2020). — BEELINE.",
    "[49] Matsumoto, H. et al. SCODE: an efficient regulatory network inference algorithm from single-cell RNA-seq during differentiation. Bioinformatics 33, 2314–2321 (2017). — SCODE.",
    "[50] Zhang, W. et al. PECA: a statistical method for reconstructing gene regulatory networks from single-cell data. Nat. Commun. 12, 5328 (2021). — PECA.",
    "[51] Schraivogel, D. et al. Targeted Perturb-seq enables genome-scale identification of cellular regulators. Nat. Biotechnol. 40, 1370–1378 (2022).",
    "[52] Adamson, B. et al. A multiplexed single-cell CRISPR screening platform enables systematic dissection of the unfolded protein response. Cell 167, 1867–1882.e21 (2016).",
    "[53] Dixit, A. et al. Perturb-Seq: dissecting molecular circuits with scalable single-cell RNA profiling of pooled genetic screens. Cell 167, 1853–1866.e17 (2016).",
    "[54] Bhatt, D. L. et al. Statins in stroke prevention. Circulation 135, 1707–1720 (2017). — Statin stroke relevance.",
    "[55] Li, Y. et al. HDAC inhibitors in neurological diseases. Nat. Rev. Drug Discov. 19, 341–359 (2020). — HDACi neuroprotection.",
    "[56] Kinney, J. W. et al. Inflammation as a central mechanism in brain injury and repair. Neurosci. Biobehav. Rev. 78, 120–134 (2017).",
    "[57] Ben-Hur, T. et al. Remyelination in the central nervous system. Nat. Rev. Neurosci. 14, 457–467 (2013). — CNS remyelination overview.",
    "[58] Nave, K. A. & Werner, H. B. Myelination of the nervous system: mechanisms and functions. Annu. Rev. Cell Dev. Biol. 30, 503–533 (2014).",
    "[59] Pauklin, S. et al. Regulatory networks in stem cell reprogramming. Nat. Rev. Genet. 15, 276–289 (2014).",
    "[60] Ramji, D. P. & Foka, P. CCAAT/enhancer-binding proteins: structure, function and regulation. Biochem. J. 365, 561–575 (2002). — C/EBP family.",
    "[61] Vicente, C. et al. The role of GATA2 in hematopoiesis and leukemogenesis. Haematologica 100, 723–733 (2015). — GATA2 biology.",
    "[62] Finzsch, M. et al. Sox9 is required for Schwann cell differentiation and myelination. Development 136, 683–693 (2009).",
    "[63] Huan, T. et al. A multi-omics approach reveals a link between peripheral blood gene expression and ischemic stroke. Stroke 50, 2331–2340 (2019).",
    "[64] Ito, H. et al. Gene expression profiling of the peri-infarct cortex following experimental stroke. J. Cereb. Blood Flow Metab. 39, 1682–1695 (2019).",
    "[65] Paschon, V. et al. Single-cell resolution of the post-stroke brain. Nat. Commun. 12, 6685 (2021).",
    "[66] Weinberg, B. H. et al. Large-scale transcriptomic profiling of stroke recovery. Cell Rep. 42, 112625 (2023).",
    "[67] Ahlmann-Eltze, C. et al. Perturbation analysis in single-cell data: a comprehensive benchmark. Nat. Methods 22, 322–331 (2025). — Perturbation benchmark.",
    "[68] Theodoris, C. V. et al. Transfer learning enables predictions in network biology. Nature 618, 616–624 (2023). — Geneformer (single-cell foundation model for network biology).",
    "[69] Kartha, V. K. et al. Functional inference of gene regulation using single-cell Perturb-seq. Nat. Genet. 55, 1339–1350 (2023).",
    "[70] Margolin, A. A. et al. ARACNE: an algorithm for the reconstruction of gene regulatory networks in a mammalian cellular context. BMC Bioinformatics 7, S7 (2006). — ARACNE (early edge-level GRN inference).",
    "[71] Edelman, G. M. & Gally, J. A. Degeneracy and complexity in biological systems. Proc. Natl Acad. Sci. USA 98, 13763–13768 (2001). — Network degeneracy.",
    "[72] Whitacre, J. M. Degeneracy: a link between evolvability, robustness and complexity in biological systems. Theor. Biol. Med. Model. 7, 6 (2010). — Degeneracy and robustness.",
    "[73] Weider, M. et al. Nfat/calcineurin signaling promotes oligodendrocyte differentiation and myelination. Nat. Commun. 12, 4240 (2021). — Sox co-regulator redundancy.",
    "[74] Huang, S. Reprogramming cell fates: reconciling rarity with robustness. BioEssays 31, 546–560 (2009). — Attractor dynamics in GRNs.",
    "[75] Raj, A. & van Oudenaarden, A. Nature, nurture, or chance: stochastic gene expression and its consequences. Cell 135, 216–226 (2008). — Single-gene expression noise.",
]
for rf in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(rf)
    set_run_font(r, size=9.5, latin="Times New Roman", ea="宋体")

# =====================================================================
# FIGURE LEGENDS
# =====================================================================
H("图注", 1)
P_rich("**图 1. 面向恢复导向型虚拟扰动的、以证据驱动的框架。** TSC-GNN 从状态条件化转录网络推断时序性重连的调控程序，并通过层级化证据框架加以解释，涵盖技术可重复性（L1）、生物学恢复（L2）、跨物种调控汇聚（L3）、转化假设生成（L4），以及情境门控的方向性因果支持（L5）。这些证据层共同逐步增强对所恢复调控程序的信心，同时明确界定框架的范围与局限。")
P_rich("**图 2. 预测基准与可解释性权衡。** （A）90 种配置下图 vs 线性的相对提升散点图。（B）示意图，阐释重新框定固定图方法贡献的“可解释性 vs 预测”权衡。")
P_rich("**图 3. 跨越卒中转换的时序重连。** 各转换显著边（汇总 q < 0.05）的 ΔW 热图：（A）sham→24 h，（B）24 h→2 d，（C）2 d→14 d，（D）sham→14 d。颜色：红 = 耦合增强，蓝 = 减弱。")
P_rich("**图 4. 跨物种调控汇聚。** （A）结构富集：SOX10→髓鞘、CEBPB→神经炎症、GATA2→神经炎症。条形：带 95% CI 的 odds ratio。（B）人类卒中血液（GSE16561）中的激活：三个恢复模块的 AUCell/ssGSEA 分数，卒中 vs 对照。")
P_rich("**图 5. 药物逆转分析。** （A）稳健 24 h 损伤特征排名靠前的 L1000 化合物。（B）置换分布（n = 20 随机特征）vs 观测计数（虚线）。（C）跨级汇聚：药物靶标与重连恢复的程序对齐。")
P_rich("**图 6. L5 三层三角验证：跨模态因果支持。** 三层证据汇聚：（i）谱系内批量 KO（L5a，阳性），（ii）LINCS 过表达（L5b，GATA2 阳性），（iii）K562 sc-CRISPRi（L5c，脱离情境空）。梯度确认因果支持依赖背景。K562 内部正对照（MYC、BCL11A）确认流程功效；GATA1 在背景内失败限定调控子可解释性。")
P_rich("**图 7. 五级证据阶梯（L1–L5）可视化概览。** 示意图展示从技术可重复性（L1）经生物学连贯性（L2）、跨物种汇聚（L3）、转化汇聚（L4）到程序级方向性因果支持（L5）的逐级上升。每一级标注其证据类型与强度。数值细节见表 1。")

# =====================================================================
# TABLES (in-text)
# =====================================================================
H("表格（文内）", 1)
P_rich("**表 1.** 五级证据阶梯（L1–L5）：设计、证据类型与强度。*[文内，§2.5；可视化概览：图 7]*")
P_rich("**表 2.** 各转换显著重连边汇总。*[文内，§3.3]*")
P_rich("**表 3.** TF 主调控因子排序的跨队列 Spearman 相关。*[文内，§3.4]*")
P_rich("**表 4.** 跨物种调控汇聚：人类同源靶的结构富集。*[文内，§3.5a]*")
P_rich("**表 5.** 人类卒中血液（GSE16561）模块激活。*[文内，§3.5b]*")
P_rich("**表 6.** Sox10 cKO（GSE269122）：靶程序方向性支持。*[文内，§3.7a]*")
P_rich("**表 7.** Cebpb het-KO（GSE273163）：靶程序方向性支持。*[文内，§3.7a]*")
P_rich("**表 8.** SigCom LINCS：特征匹配结果。*[文内，§3.7b]*")
P_rich("**表 9.** K562 sc-CRISPR：调控子响应结果。*[文内，§3.7c]*")
P_rich("**表 10.** 跨模态扰动支持图谱。*[文内，§3.8.2]*")
P_rich("**表 11.** K562 正对照调控子响应。*[文内，§3.8.2]*")

# =====================================================================
# SUPPLEMENTARY MATERIALS
# =====================================================================
H("补充材料（拟）", 1)
B("补充图 S1：PC 组成校正：全部显著边的原始 vs 校正 ΔW")
B("补充图 S2：跨越卒中转换的细胞间通讯重塑")
B("补充图 S3：随机图重连负对照")
B("补充表 S1：每个 TF 的 DoRothEA TF→靶边数（小鼠、人类）")
B("补充表 S2：完整 L1000CDS2 药物逆转列表")
B("补充表 S3：K562 sc-CRISPR 完整调控子响应排名")
B("补充代码 S1：TSC-GNN 流程（conda 环境 + 脚本）")

doc.save(OUT)
print("SAVED:", OUT)
